

"""
===============================================================================
 Created on Feb 15, 2020

 @author: Ramin Mehdizad Tekiyeh

 This code is written in Python 3.7.4 , Spyder 3.3.6
===============================================================================
"""

#==============================================================================
# This module contains word docx codes
#==============================================================================


#==============================================================================
# imports
#==============================================================================
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import ns
from docx.shared import Cm, Mm


#==============================================================================
# importing module codes
#==============================================================================
import ModVar as Var


#==============================================================================
# some text information
#==============================================================================
Report_File_Name='Auto_Created_Report.docx'

author='Ramin Mehdizad'
Header_ttl='Music hit Database'
date='Apr., 17, 2020'

pic_width=2.2

sep_line='_' * 54

FigSubttl_txt='Fig. 1: ...'

sec1_ttl='Introduction'
sec1_txt=('This code analyzes Music-Hit database and Makes predictions to' 
          'the test data. Then saves it into a *.csv file. ...') * 6
                        
sec2_ttl='Dataset Description'
sec2_txt='Dataset consists of ...'

sec3_ttl='Preprocessing'
sec3_txt='In the preprocessing level, ...'

sec4_ttl='Algorithm Implementation'
sec4_txt=('Four methods of random forest, KNN, SVN and Logistic'
       'Regression is implementes with various parameters in order to'
        'find the optimum ...')
        
sec5_ttl='Part 1: Random Forest'
sec5_txt=('The results of analysis for implementation of Random Forest' 
               'is presented in this section. ...')
    
sec6_ttl='Part 2: KNN'
sec6_txt=('The results of analysis for implementation of KNN is presented '
             'in this section. ...')
    
sec7_ttl='Part 3: SVM'
sec7_txt=('The results of analysis for implementation of SVM is presented' 
            'in this section. ...')
    
sec8_ttl='Part 4: Logistic Regression'
sec8_txt=('The results of analysis for implementation of '
            'Logistic Regression is presented in this section. ...')

sec9_ttl='Results'
sec9_txt=('In the following table, ....')

sec10_ttl='Conclusion'
sec10_txt=('By analyzing the results of the four classifier tyoes on the '
          'dataset, the following conclusions can be made: ...')


#==============================================================================
# this function creates page number 
#==============================================================================
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    

#==============================================================================
# this section creates paragraph styles
#============================================================================== 
def Crt_Styles():
    global styles
    
    # creating simple paragrach styles
    styles = document.styles
    style = styles.add_style('TxtStyle', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(0,0,0)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Pt(12)
    paragraph_format.first_line_indent = Pt(12)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(12)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = 1
    
    # creating figure subtitle styles
    styles = document.styles
    style = styles.add_style('FigSubttlStyle', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    font.italic = False
    font.bold=True
    font.underline = False
    font.color.rgb = RGBColor(0,0,0)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(12)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.line_spacing = 1

    # creating header1 styles
    styles = document.styles
    style = styles.add_style('Header1', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.italic = False
    font.bold=True
    font.underline = False
    font.color.rgb = RGBColor(0,0,255)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # creating page header style
    styles = document.styles
    style = styles.add_style('PageHeader', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.italic = True
    font.bold=True
    font.underline = False
    font.color.rgb = RGBColor(0,0,255)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # creating separator line format
    styles = document.styles
    style = styles.add_style('SepLine', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(16)
    font.italic = False
    font.bold=True
    font.underline = False
    font.color.rgb = RGBColor(255,192,0)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # creating picture paragraph format
    styles = document.styles
    style = styles.add_style('Pic', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
#==============================================================================
# this function creates report file 
#==============================================================================
def Rep():
    global document
    
    #-------------------------preliminary settings-----------------------------
    # create document
    document = Document()
    
    # create styles
    Crt_Styles()
    
    # add page No.
    add_page_number(document.sections[0].footer.paragraphs[0].add_run())

    # add page header
    section = document.sections[0]
    header = section.header
    paragraph_curr = header.paragraphs[0]
    paragraph_curr.text = "{}\t\t{}\t\t{}".format(author,Header_ttl,date)
    paragraph_curr.style = document.styles['PageHeader']
    paragraph_curr = header.add_paragraph(sep_line)
    paragraph_curr.style = document.styles['SepLine']

    # page size
    section.page_height=Mm(297)
    section.page_width=Mm(210)
    section.left_margin=Mm(25.4)
    section.right_margin=Mm(25.4)
    section.top_margin=Mm(25.4)
    section.bottom_margin=Mm(25.4)
    section.header_distance=Mm(12.7)
    section.footer_distance=Mm(12.7)
    #--------------------------------------------------------------------------



    #-------------------------writing text and plots---------------------------
    # writing introduction
    paragraph_curr = document.add_paragraph(sec1_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec1_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    
    # writing dataset description
    paragraph_curr = document.add_paragraph(sec2_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec2_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    
    # writing preprocessing
    paragraph_curr = document.add_paragraph(sec3_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec3_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    
    # writing Algorithm Implementation
    paragraph_curr = document.add_paragraph(sec4_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec4_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    # inserting pictures of features importance
    paragraph_curr = document.add_paragraph()
    Run=paragraph_curr.add_run()
    paragraph_curr.style = document.styles['Pic']
    for i in range(2):
        j=i+1
        pic_name='Imp_' + str(j) + '.jpg'
        Run.add_picture(pic_name, width=Inches(pic_width))
    # picture caption
    paragraph_curr = document.add_paragraph(FigSubttl_txt)
    paragraph_curr.style = document.styles['FigSubttlStyle']
    # inserting pictures of PCA
    paragraph_curr = document.add_paragraph()
    Run=paragraph_curr.add_run()
    paragraph_curr.style = document.styles['Pic']
    for i in range(2):
        j=i+1
        pic_name='PCA_' + str(j) + '.jpg'
        Run.add_picture(pic_name, width=Inches(pic_width))
    # picture caption
    paragraph_curr = document.add_paragraph(FigSubttl_txt)
    paragraph_curr.style = document.styles['FigSubttlStyle']
    
    # writing  Part 1: RF
    paragraph_curr = document.add_paragraph(sep_line)
    paragraph_curr.style = document.styles['SepLine']
    paragraph_curr = document.add_paragraph(sec5_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec5_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    # inserting pictures
    paragraph_curr = document.add_paragraph()
    Run=paragraph_curr.add_run()
    paragraph_curr.style = document.styles['Pic']
    for i in range(4):
        j=i+1
        pic_name='RF_' + str(j) + '.jpg'
        Run.add_picture(pic_name, width=Inches(pic_width))
    # picture caption
    paragraph_curr = document.add_paragraph(FigSubttl_txt)
    paragraph_curr.style = document.styles['FigSubttlStyle']
    
    # writing  Part 2: knn
    paragraph_curr = document.add_paragraph(sep_line)
    paragraph_curr.style = document.styles['SepLine']
    paragraph_curr = document.add_paragraph(sec6_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec6_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    # inserting pictures
    paragraph_curr = document.add_paragraph()
    Run=paragraph_curr.add_run()
    paragraph_curr.style = document.styles['Pic']
    for i in range(4):
        j=i+1
        pic_name='KNN_' + str(j) + '.jpg'
        Run.add_picture(pic_name, width=Inches(pic_width))
    # picture caption
    paragraph_curr = document.add_paragraph(FigSubttl_txt)
    paragraph_curr.style = document.styles['FigSubttlStyle']
    
    # writing  Part 3: svm
    paragraph_curr = document.add_paragraph(sep_line)
    paragraph_curr.style = document.styles['SepLine']
    paragraph_curr = document.add_paragraph(sec7_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec7_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    # inserting pictures    
    paragraph_curr = document.add_paragraph()
    Run=paragraph_curr.add_run()
    paragraph_curr.style = document.styles['Pic']
    for i in range(4):
        j=i+1
        pic_name='SVM_' + str(j) + '.jpg'
        Run.add_picture(pic_name, width=Inches(pic_width))
    # picture caption
    paragraph_curr = document.add_paragraph(FigSubttl_txt)
    paragraph_curr.style = document.styles['FigSubttlStyle']
    
    # writing  Part 4: log reg
    paragraph_curr = document.add_paragraph(sep_line)
    paragraph_curr.style = document.styles['SepLine']
    paragraph_curr = document.add_paragraph(sec8_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec8_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    # inserting pictures    
    paragraph_curr = document.add_paragraph()
    Run=paragraph_curr.add_run()
    paragraph_curr.style = document.styles['Pic']
    for i in range(4):
        j=i+1
        pic_name='LR_' + str(j) + '.jpg'
        Run.add_picture (pic_name, width=Inches(pic_width))
    # picture caption
    paragraph_curr = document.add_paragraph(FigSubttl_txt)
    paragraph_curr.style = document.styles['FigSubttlStyle']
    
    # writing results
    paragraph_curr = document.add_paragraph(sep_line)
    paragraph_curr.style = document.styles['SepLine']
    paragraph_curr = document.add_paragraph(sec9_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec9_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    
    # insert tables for time of run
    table1 = document.add_table(rows=2, cols=len(Var.Time_List)+1, 
                                style='Medium Grid 3 Accent 3')
    table1.allow_autofit = True
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
   
    # set header
    hdr=['','RF','KNN','SVM','LR','Total']
    hdr_cells = table1.rows[0].cells
    for i,t in enumerate(hdr):
        hdr_cells[i].text = t
        
    # set row names
    rowName='Run Time (s)'
    record_cells = table1.rows[1].cells
    record_cells[0].text = rowName
    # writing records
    records=list()
    for i,t in enumerate(Var.Time_List):
        records.append(str(t))
    record_cells = table1.rows[1].cells
    for i,t in enumerate(records):
        record_cells[i+1].text = t
    
    # writing conclusion
    paragraph_curr = document.add_paragraph(sep_line)
    paragraph_curr.style = document.styles['SepLine']
    paragraph_curr = document.add_paragraph(sec10_ttl)
    paragraph_curr.style = document.styles['Header1']
    paragraph_curr = document.add_paragraph(sec10_txt)
    paragraph_curr.style = document.styles['TxtStyle']
    #--------------------------------------------------------------------------


    #----------------------------finalizing------------------------------------
    # save report file
    document.save(Report_File_Name)
    # after use we delete created styles
    styles['TxtStyle'].delete()
    styles['Header1'].delete()
    styles['PageHeader'].delete()
    styles['SepLine'].delete()
    styles['Pic'].delete()
    styles['FigSubttlStyle'].delete()
    #--------------------------------------------------------------------------







